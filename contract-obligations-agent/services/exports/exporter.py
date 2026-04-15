from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document


def export_json(result: dict, path: Path) -> Path:
    path.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def export_markdown(result: dict, path: Path) -> Path:
    lines = [f"# {result['filename']}", "", "## Executive summary", result["summary"]["executive_summary"], "", "## Key points"]
    lines.extend(f"- {point}" for point in result.get("summary", {}).get("key_points", []))
    lines.extend(["", "## Clauses"])
    for clause in result.get("clauses", []):
        lines.append(f"- [{clause.get('clause_type', '')}] {clause.get('title', '')}")
    lines.extend(["", "## Obligations"])
    lines.extend(f"- {item['description']}" for item in result.get("obligations", []))
    if result.get("dates"):
        lines.extend(["", "## Dates"])
        lines.extend(f"- {item.get('date', '')}: {item.get('context', '')}" for item in result["dates"])
    if result.get("retrieval_hits"):
        lines.extend(["", "## Evidence hits"])
        for hit in result["retrieval_hits"]:
            lines.append(f"- [{hit.get('rank', '?')}] {hit.get('source_label', '')}: {hit.get('source_excerpt', '')}")
    if result.get("comparison"):
        lines.extend(["", "## Checklist comparison"])
        comparison = result["comparison"]
        lines.append(f"- Status: {comparison.get('status', '')}")
        for key in ("matched", "missing", "unverifiable"):
            items = comparison.get(key, [])
            if items:
                lines.append(f"- {key.title()}: {', '.join(map(str, items))}")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def export_csv(result: dict, path: Path) -> Path:
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=["obligation_id", "description", "responsible_party", "due_date", "dependency", "status", "confidence"])
        writer.writeheader()
        for item in result.get("obligations", []):
            writer.writerow({k: item.get(k, "") for k in writer.fieldnames})
    return path


def export_docx(result: dict, path: Path) -> Path:
    doc = Document()
    doc.add_heading(result["filename"], level=1)
    doc.add_heading("Executive summary", level=2)
    doc.add_paragraph(result["summary"]["executive_summary"])
    for point in result.get("summary", {}).get("key_points", []):
        doc.add_paragraph(point, style="List Bullet")
    if result.get("comparison"):
        doc.add_heading("Checklist comparison", level=2)
        comparison = result["comparison"]
        doc.add_paragraph(f"Status: {comparison.get('status', '')}")
        for key in ("matched", "missing", "unverifiable"):
            items = comparison.get(key, [])
            if items:
                doc.add_paragraph(f"{key.title()}: {', '.join(map(str, items))}")
    doc.add_heading("Obligations", level=2)
    for item in result.get("obligations", []):
        doc.add_paragraph(item["description"], style="List Bullet")
    if result.get("retrieval_hits"):
        doc.add_heading("Evidence hits", level=2)
        for hit in result["retrieval_hits"]:
            doc.add_paragraph(f"[{hit.get('rank', '?')}] {hit.get('source_label', '')}: {hit.get('source_excerpt', '')}", style="List Bullet")
    doc.save(path)
    return path


def export_bundle(result: dict, directory: Path) -> dict[str, Path]:
    directory.mkdir(parents=True, exist_ok=True)
    return {
        "json": export_json(result, directory / "analysis.json"),
        "markdown": export_markdown(result, directory / "analysis.md"),
        "csv": export_csv(result, directory / "obligations.csv"),
        "docx": export_docx(result, directory / "analysis.docx"),
    }
