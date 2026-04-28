"""
domain/exporters.py
-------------------
Motor de exportacin profesional a formato Microsoft Word (.docx).
Inspirado en la Fase 16 del proyecto p01.
"""

from __future__ import annotations
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config.settings import REPORTS_DIR

def generar_nombre_archivo(base: str = "Briefing_Ejecutivo") -> Path:
    """Genera una ruta nica con timestamp para evitar sobreescritura."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return REPORTS_DIR / f"{base}_{timestamp}.docx"

def anadir_portada(doc: Document, titulo: str, subtitulo: str):
    """Aade una portada profesional con estilo ejecutivo."""
    # Ttulo principal
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run(titulo)
    run_t.bold = True
    run_t.font.size = Pt(24)
    run_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) # Azul corporativo

    # Subttulo
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run(subtitulo)
    run_s.font.size = Pt(14)
    run_s.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph("\n" * 2)
    
    # Metadatos de generacin
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Fecha de generacin: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_page_break()

def exportar_a_word(mensajes: list[dict], nombre_pdf: str = "N/A") -> Path:
    """
    Toma el historial del chat y genera un informe estructurado.
    Separa el contenido del usuario y las respuestas finales de la IA.
    """
    doc = Document()
    
    # 1. Portada
    anadir_portada(
        doc, 
        "INFORME DE INTELIGENCIA ESTRATGICA", 
        f"Analisis basado en: {nombre_pdf}"
    )
    
    # 2. Resumen del Analisis
    doc.add_heading("1. Resumen de la Consulta", level=1)
    
    for msg in mensajes:
        role = "USUARIO" if msg["role"] == "user" else "ASISTENTE"
        
        # Saltamos los mensajes tcunicos internos si los hubiera
        if msg["role"] not in ["user", "assistant"]:
            continue
            
        p = doc.add_paragraph()
        run_role = p.add_run(f"[{role}]: ")
        run_role.bold = True
        p.add_run(msg["content"])
        doc.add_paragraph() # Espacio extra

    # 3. Pie de pagina metodologico
    doc.add_page_break()
    doc.add_heading("Nota Metodologica", level=2)
    p_nota = doc.add_paragraph(
        "Este informe ha sido generado mediante un Asistente Ejecutivo IA con capacidades "
        "de busqueda web en tiempo real y analisis de documentos locales. "
        "La veracidad de los datos externos depende de las fuentes consultadas."
    )
    p_nota.runs[0].font.italic = True
    p_nota.runs[0].font.size = Pt(8)

    # 4. Guardado
    output_path = generar_nombre_archivo()
    doc.save(output_path)
    return output_path