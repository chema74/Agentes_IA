from __future__ import annotations

import csv
import io
import json
import zipfile
from pathlib import Path
from typing import Any

from docx import Document


def _as_data(value: Any) -> Any:
    if hasattr(value, "model_dump"):
        return value.model_dump(mode="json")
    if isinstance(value, dict):
        return {key: _as_data(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_as_data(item) for item in value]
    return value


def _write_json(data: Any, path: Path) -> Path:
    path.write_text(json.dumps(_as_data(data), indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def _write_markdown(payload: dict[str, Any], path: Path) -> Path:
    controls = payload.get("controls", [])
    evidence = payload.get("evidence", [])
    findings = payload.get("findings", [])
    scope = payload.get("scope", {})
    lines = [
        f"# {payload.get('package_name', 'Audit package')}",
        "",
        f"- Scope: {scope.get('name', scope.get('id', 'unknown'))}",
        f"- Controls: {len(controls)}",
        f"- Evidence items: {len(evidence)}",
        f"- Findings: {len(findings)}",
        "",
        "## Controls",
    ]
    for control in controls:
        lines.append(f"- {control.get('code', 'N/A')} {control.get('name', '')}".strip())
    if not controls:
        lines.append("- None")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def _write_csv(payload: dict[str, Any], path: Path) -> Path:
    controls = payload.get("controls", [])
    fieldnames = ["id", "code", "name", "category", "criticality", "status"]
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for control in controls:
            writer.writerow({key: control.get(key, "") for key in fieldnames})
    return path


def _write_docx(payload: dict[str, Any], path: Path) -> Path:
    doc = Document()
    doc.add_heading(payload.get("package_name", "Audit package"), level=1)
    scope = payload.get("scope", {})
    doc.add_paragraph(f"Scope: {scope.get('name', scope.get('id', 'unknown'))}")
    doc.add_paragraph(f"Controls: {len(payload.get('controls', []))}")
    doc.add_paragraph(f"Evidence: {len(payload.get('evidence', []))}")
    doc.add_paragraph(f"Findings: {len(payload.get('findings', []))}")
    doc.save(path)
    return path


def build_audit_package_zip(payload: Any) -> bytes:
    data = _as_data(payload)
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("cover.md", _render_cover_md(data))
        archive.writestr("controls.csv", _render_controls_csv(data))
        archive.writestr("audit_package.docx", _render_docx_bytes(data))
        archive.writestr("package.json", json.dumps(data, indent=2, ensure_ascii=False))
    return buffer.getvalue()


def _render_cover_md(payload: dict[str, Any]) -> str:
    scope = payload.get("scope", {})
    return "\n".join(
        [
            f"# {payload.get('package_name', 'Audit package')}",
            "",
            f"Scope: {scope.get('name', scope.get('id', 'unknown'))}",
            f"Controls: {len(payload.get('controls', []))}",
            f"Evidence items: {len(payload.get('evidence', []))}",
            f"Findings: {len(payload.get('findings', []))}",
            "",
        ]
    )


def _render_controls_csv(payload: dict[str, Any]) -> str:
    fieldnames = ["id", "code", "name", "category", "criticality", "status"]
    buffer = io.StringIO()
    writer = csv.DictWriter(buffer, fieldnames=fieldnames)
    writer.writeheader()
    for control in payload.get("controls", []):
        writer.writerow({key: control.get(key, "") for key in fieldnames})
    return buffer.getvalue()


def _render_docx_bytes(payload: dict[str, Any]) -> bytes:
    with io.BytesIO() as buffer:
        doc = Document()
        doc.add_heading(payload.get("package_name", "Audit package"), level=1)
        scope = payload.get("scope", {})
        doc.add_paragraph(f"Scope: {scope.get('name', scope.get('id', 'unknown'))}")
        doc.add_paragraph(f"Controls: {len(payload.get('controls', []))}")
        doc.add_paragraph(f"Evidence: {len(payload.get('evidence', []))}")
        doc.add_paragraph(f"Findings: {len(payload.get('findings', []))}")
        doc.save(buffer)
        return buffer.getvalue()
