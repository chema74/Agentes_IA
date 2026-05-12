from __future__ import annotations

import hashlib
import re

from connectors.email.loader import parse_eml
from connectors.files.loaders import FileLoadError, LoadedFile
from domain.contracts.models import DocumentType, ParsedDocument, TextChunk


WHITESPACE_RE = re.compile(r"[ \t]+")
NEWLINES_RE = re.compile(r"\n{3,}")
HYPHEN_BREAK_RE = re.compile(r"-\n(?=\w)")
SPACE_AROUND_NEWLINES_RE = re.compile(r"[ \t]*\n[ \t]*")


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = HYPHEN_BREAK_RE.sub("", text)
    text = WHITESPACE_RE.sub(" ", text)
    text = SPACE_AROUND_NEWLINES_RE.sub("\n", text)
    text = NEWLINES_RE.sub("\n\n", text)
    return text.strip()


def _chunk_text(text: str, source_label: str, size: int = 1200) -> list[TextChunk]:
    out: list[TextChunk] = []
    for idx, start in enumerate(range(0, len(text), size), start=1):
        piece = text[start : start + size].strip()
        if piece:
            out.append(TextChunk(chunk_id=f"{source_label}-chunk-{idx}", text=piece, source_label=source_label, start_char=start, end_char=min(start + size, len(text))))
    return out


def detect_document_type(filename: str, suffix: str, metadata: dict | None = None) -> DocumentType:
    suffix = suffix.lower()
    name = filename.lower()
    if suffix == ".eml":
        return DocumentType.email
    if suffix == ".docx":
        if "template" in name or (metadata or {}).get("kind") == "template":
            return DocumentType.template
        if "annex" in name or "anexo" in name:
            return DocumentType.annex
        return DocumentType.contract
    if suffix == ".pdf":
        return DocumentType.contract
    return DocumentType.unknown


def _parse_docx(content: bytes) -> tuple[str, dict]:
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs), {"paragraphs": len(paragraphs), "tables": len(doc.tables)}


def _parse_eml(content: bytes) -> tuple[str, dict]:
    parsed = parse_eml(content)
    parts = [f"{k.upper()}: {v}" for k, v in parsed.items() if k != "body" and v]
    if parsed.get("body"):
        parts.append(parsed["body"])
    return "\n\n".join(parts), {k: parsed.get(k, "") for k in ("subject", "from", "to", "date")}


def _parse_pdf(content: bytes) -> tuple[str, dict]:
    try:
        from pypdf import PdfReader
        from io import BytesIO
    except Exception as exc:
        raise FileLoadError("PDF support requires pypdf in the environment.") from exc
    reader = PdfReader(BytesIO(content))
    text = []
    for page in reader.pages:
        text.append(page.extract_text() or "")
    return "\n\n".join(text), {"pages": len(reader.pages)}


def parse_loaded_file(loaded: LoadedFile) -> ParsedDocument:
    if not loaded.content:
        raise FileLoadError("Cannot parse empty file.")
    if loaded.suffix == ".docx":
        raw_text, meta = _parse_docx(loaded.content)
    elif loaded.suffix == ".eml":
        raw_text, meta = _parse_eml(loaded.content)
    elif loaded.suffix == ".pdf":
        raw_text, meta = _parse_pdf(loaded.content)
    else:
        raw_text, meta = loaded.content.decode("utf-8", errors="ignore"), {}
    normalized = normalize_text(raw_text)
    if not normalized:
        raise FileLoadError(f"No extractable text found in {loaded.filename}.")
    doc_id = hashlib.sha256(loaded.content).hexdigest()[:24]
    return ParsedDocument(
        document_id=doc_id,
        filename=loaded.filename,
        document_type=detect_document_type(loaded.filename, loaded.suffix, meta),
        metadata=meta,
        raw_text=raw_text,
        normalized_text=normalized,
        chunks=_chunk_text(normalized, loaded.filename),
    )
