from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

from docx import Document


def _as_data(value: Any) -> Any:
    if hasattr(value, "model_dump"):
        return value.model_dump(mode="json")
    if isinstance(value, dict):
        return {key: _as_data(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_as_data(item) for item in value]
    return value


def _payload(data: Any) -> dict[str, Any]:
    normalized = _as_data(data)
    return normalized if isinstance(normalized, dict) else {"value": normalized}


def export_json(data: Any, path: str | Path) -> Path:
    path = Path(path)
    path.write_text(json.dumps(_as_data(data), indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def export_markdown(data: Any, path: str | Path) -> Path:
    path = Path(path)
    payload = _payload(data)
    summary = payload.get("summary", {})
    obligations = payload.get("obligations", [])
    lines = [
        f"# {payload.get('filename', 'Contract analysis')}",
        "",
        f"## Summary",
        f"- {summary.get('executive_summary', '')}",
    ]
    key_points = summary.get("key_points", [])
    if key_points:
        lines.append("")
        lines.append("### Key points")
        lines.extend(f"- {point}" for point in key_points)
    if obligations:
        lines.append("")
        lines.append("## Obligations")
        for item in obligations:
            lines.append(f"- {item.get('obligation_id', 'n/a')}: {item.get('description', '')}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def export_csv(data: Any, path: str | Path) -> Path:
    path = Path(path)
    payload = _payload(data)
    obligations = payload.get("obligations", [])
    fieldnames = ["obligation_id", "description", "responsible_party", "due_date", "dependency", "status", "confidence"]
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for item in obligations:
            writer.writerow({key: item.get(key, "") for key in fieldnames})
    return path


def export_docx(data: Any, path: str | Path) -> Path:
    path = Path(path)
    payload = _payload(data)
    summary = payload.get("summary", {})
    obligations = payload.get("obligations", [])
    doc = Document()
    doc.add_heading(payload.get("filename", "Contract analysis"), level=1)
    doc.add_paragraph(summary.get("executive_summary", ""))
    if summary.get("key_points"):
        doc.add_heading("Key points", level=2)
        for point in summary["key_points"]:
            doc.add_paragraph(point, style="List Bullet")
    if obligations:
        doc.add_heading("Obligations", level=2)
        table = doc.add_table(rows=1, cols=4)
        headers = ["ID", "Description", "Party", "Due date"]
        for cell, header in zip(table.rows[0].cells, headers, strict=False):
            cell.text = header
        for item in obligations:
            row = table.add_row().cells
            row[0].text = str(item.get("obligation_id", ""))
            row[1].text = str(item.get("description", ""))
            row[2].text = str(item.get("responsible_party", ""))
            row[3].text = str(item.get("due_date", ""))
    doc.save(path)
    return path


def export_bundle(data: Any, directory: str | Path) -> dict[str, Path]:
    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    payload = _payload(data)
    return {
        "json": export_json(payload, directory / "result.json"),
        "markdown": export_markdown(payload, directory / "result.md"),
        "csv": export_csv(payload, directory / "result.csv"),
        "docx": export_docx(payload, directory / "result.docx"),
    }
