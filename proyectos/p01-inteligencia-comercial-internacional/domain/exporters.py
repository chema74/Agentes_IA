from __future__ import annotations

import csv
import json
import re
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from domain.schemas import RankingResult


OUTPUT_DIR = Path("outputs")
REPORTS_OUTPUT_DIR = OUTPUT_DIR / "reports"
RANKINGS_OUTPUT_DIR = OUTPUT_DIR / "rankings"


def ensure_output_dirs() -> None:
    """
    Crea los directorios necesarios para exportaciones.
    """
    REPORTS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    RANKINGS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def slugify_text(text: str) -> str:
    """
    Convierte un texto en una versión segura para nombres de archivo.
    """
    normalized = unicodedata.normalize("NFKD", text.strip())
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_text = ascii_text.lower().replace(" ", "_")
    ascii_text = re.sub(r"[^a-z0-9_\-]", "", ascii_text)
    return ascii_text or "sin_valor"


def prettify_dimension_name(name: str) -> str:
    """
    Convierte una clave técnica a una etiqueta legible.
    """
    labels = {
        "riesgo_politico": "Riesgo político",
        "riesgo_comercial": "Riesgo comercial",
        "estabilidad_economica": "Estabilidad económica",
        "riesgo_regulatorio": "Riesgo regulatorio",
        "riesgo_geopolitico": "Riesgo geopolítico",
        "riesgo_operativo": "Riesgo operativo",
        "ajuste_sectorial": "Ajuste sectorial",
        "oportunidad_sectorial": "Oportunidad sectorial",
        "politico_institucional": "Riesgo político e institucional",
        "macroeconomico": "Riesgo macroeconómico",
        "regulatorio_cumplimiento": "Riesgo regulatorio y de cumplimiento",
        "geopolitico": "Riesgo geopolítico",
        "comercial_operativo": "Riesgo comercial y operativo",
        "sectorial": "Riesgo sectorial",
    }
    return labels.get(name, name.replace("_", " ").capitalize())


def build_filename(prefix: str, pais_a: str, pais_b: str | None = None) -> Path:
    """
    Construye un nombre base de archivo para comparación de países.
    """
    ensure_output_dirs()

    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    pais_a_slug = slugify_text(pais_a)

    if pais_b:
        pais_b_slug = slugify_text(pais_b)
        filename = f"{prefix}_{timestamp}_{pais_a_slug}_{pais_b_slug}"
    else:
        filename = f"{prefix}_{timestamp}_{pais_a_slug}"

    return REPORTS_OUTPUT_DIR / filename


def _normalize_base_path(filename_without_ext: str | Path) -> Path:
    """
    Normaliza una ruta base sin extensión.
    """
    path = Path(filename_without_ext)
    if path.suffix:
        path = path.with_suffix("")
    return path


def export_to_json(data: dict, filename_without_ext: str | Path) -> Path:
    """
    Exporta un diccionario a JSON.
    """
    ensure_output_dirs()
    base_path = _normalize_base_path(filename_without_ext)
    output_path = base_path.with_suffix(".json")

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

    return output_path


def _append_country_markdown(lines: list[str], country_name: str, result: dict) -> None:
    """
    Añade al informe Markdown el bloque completo de un país.
    """
    lines.append(f"## {country_name}")
    lines.append("")

    resumen = result.get("resumen_ejecutivo", "Sin resumen ejecutivo disponible.")
    lines.append("### Resumen ejecutivo")
    lines.append(resumen)
    lines.append("")

    scores = result.get("scores", {})
    lines.append("### Scores")
    for key, value in scores.items():
        lines.append(f"- **{prettify_dimension_name(key)}:** {value}")
    lines.append("")

    justificaciones = result.get("justificacion_scores", {})
    lines.append("### Justificación de scores")
    if justificaciones:
        for dimension, motivos in justificaciones.items():
            lines.append(f"- **{prettify_dimension_name(dimension)}:**")
            for motivo in motivos:
                lines.append(f"  - {motivo}")
    else:
        lines.append("- Sin justificaciones registradas.")
    lines.append("")

    fuentes = result.get("fuentes", [])
    lines.append("### Fuentes")
    if fuentes:
        for fuente in fuentes:
            categoria = fuente.get("categoria", "sin_categoria")
            titulo = fuente.get("titulo", "Sin título")
            url = fuente.get("url", "")
            resumen_fuente = fuente.get("resumen", "")

            lines.append(f"- **[{categoria.upper()}] {titulo}**")
            if url:
                lines.append(f"  - URL: {url}")
            if resumen_fuente:
                lines.append(f"  - Resumen: {resumen_fuente}")
    else:
        lines.append("- No se registraron fuentes.")
    lines.append("")


def build_markdown_report(
    pais_a: str,
    resultado_a: dict,
    pais_b: str,
    resultado_b: dict,
    comparativa_scores: dict,
    comparacion_narrativa: str,
    sector: str,
    tipo_empresa: str,
) -> str:
    """
    Construye un informe Markdown para la comparación entre dos países.
    """
    lines: list[str] = []

    lines.append("# Informe comparativo de riesgo país")
    lines.append("")
    lines.append("## Metadatos")
    lines.append(f"- Fecha de generación: {datetime.now().isoformat(timespec='seconds')}")
    lines.append(f"- Sector: {sector}")
    lines.append(f"- Tipo de empresa: {tipo_empresa}")
    lines.append(f"- País A: {pais_a}")
    lines.append(f"- País B: {pais_b}")
    lines.append("")

    _append_country_markdown(lines, pais_a, resultado_a)
    _append_country_markdown(lines, pais_b, resultado_b)

    lines.append("## Comparativa estructurada de scores")
    lines.append("")
    for clave, valor in comparativa_scores.items():
        lines.append(f"- **{clave}:** {valor}")
    lines.append("")

    lines.append("## Comparación narrativa final")
    lines.append("")
    lines.append(comparacion_narrativa or "Sin comparación narrativa disponible.")
    lines.append("")

    return "\n".join(lines)


def export_to_markdown(markdown_text: str, filename_without_ext: str | Path) -> Path:
    """
    Exporta texto Markdown a archivo .md
    """
    ensure_output_dirs()
    base_path = _normalize_base_path(filename_without_ext)
    output_path = base_path.with_suffix(".md")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(markdown_text)

    return output_path


def _append_country_docx(document: Document, country_name: str, result: dict) -> None:
    """
    Añade al DOCX el bloque completo de un país.
    """
    document.add_heading(country_name, level=1)

    document.add_heading("Resumen ejecutivo", level=2)
    document.add_paragraph(result.get("resumen_ejecutivo", "Sin resumen ejecutivo disponible."))

    document.add_heading("Scores", level=2)
    scores = result.get("scores", {})
    if scores:
        for key, value in scores.items():
            document.add_paragraph(f"{prettify_dimension_name(key)}: {value}")
    else:
        document.add_paragraph("Sin scores registrados.")

    document.add_heading("Justificación de scores", level=2)
    justificaciones = result.get("justificacion_scores", {})
    if justificaciones:
        for dimension, motivos in justificaciones.items():
            document.add_paragraph(prettify_dimension_name(dimension))
            for motivo in motivos:
                p = document.add_paragraph(style="List Bullet")
                p.add_run(motivo)
    else:
        document.add_paragraph("Sin justificaciones registradas.")

    document.add_heading("Fuentes", level=2)
    fuentes = result.get("fuentes", [])
    if fuentes:
        for fuente in fuentes:
            categoria = fuente.get("categoria", "sin_categoria")
            titulo = fuente.get("titulo", "Sin título")
            url = fuente.get("url", "")
            resumen_fuente = fuente.get("resumen", "")

            p = document.add_paragraph(style="List Bullet")
            p.add_run(f"[{categoria.upper()}] {titulo}")
            if url:
                p.add_run(f" — {url}")
            if resumen_fuente:
                document.add_paragraph(resumen_fuente)
    else:
        document.add_paragraph("No se registraron fuentes.")


def export_to_docx(
    pais_a: str,
    resultado_a: dict,
    pais_b: str,
    resultado_b: dict,
    comparativa_scores: dict,
    comparacion_narrativa: str,
    sector: str,
    tipo_empresa: str,
    filename_without_ext: str | Path,
) -> Path:
    """
    Exporta la comparación entre dos países a DOCX.
    """
    ensure_output_dirs()
    base_path = _normalize_base_path(filename_without_ext)
    output_path = base_path.with_suffix(".docx")

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    document.add_heading("Informe comparativo de riesgo país", level=0)
    document.add_paragraph(f"Fecha de generación: {datetime.now().isoformat(timespec='seconds')}")
    document.add_paragraph(f"Sector: {sector}")
    document.add_paragraph(f"Tipo de empresa: {tipo_empresa}")
    document.add_paragraph(f"País A: {pais_a}")
    document.add_paragraph(f"País B: {pais_b}")

    _append_country_docx(document, pais_a, resultado_a)
    _append_country_docx(document, pais_b, resultado_b)

    document.add_heading("Comparativa estructurada de scores", level=1)
    for clave, valor in comparativa_scores.items():
        document.add_paragraph(f"{clave}: {valor}")

    document.add_heading("Comparación narrativa final", level=1)
    document.add_paragraph(comparacion_narrativa or "Sin comparación narrativa disponible.")

    document.save(output_path)
    return output_path


def ensure_rankings_output_dir() -> Path:
    """
    Garantiza la existencia del directorio de salida para rankings.
    """
    ensure_output_dirs()
    return RANKINGS_OUTPUT_DIR


def build_ranking_filename(
    sector: str,
    company_type: str,
    extension: str,
    prefix: str = "ranking",
    run_id: str | None = None,
) -> Path:
    """
    Construye el nombre de archivo del ranking con timestamp coherente.
    """
    output_dir = ensure_rankings_output_dir()

    if run_id is None:
        run_id = datetime.now().strftime("%Y-%m-%d_%H%M%S")

    sector_slug = slugify_text(sector)
    company_type_slug = slugify_text(company_type)

    filename = f"{prefix}_{run_id}_{sector_slug}_{company_type_slug}.{extension}"
    return output_dir / filename


def _collect_dimension_keys(ranking_result: RankingResult) -> list[str]:
    """
    Obtiene todas las claves de dimensiones presentes en el ranking.
    """
    keys: set[str] = set()
    for item in ranking_result.ranking:
        keys.update(item.dimension_scores.keys())
    return sorted(keys)


def flatten_ranking_for_csv(ranking_result: RankingResult) -> List[Dict[str, Any]]:
    """
    Convierte la estructura jerárquica del ranking en filas planas para CSV.
    """
    rows: List[Dict[str, Any]] = []
    dimension_keys = _collect_dimension_keys(ranking_result)

    for item in ranking_result.ranking:
        row: Dict[str, Any] = {
            "position": item.position,
            "country": item.country,
            "score_total": item.score_total,
            "executive_summary": item.executive_summary,
            "sources_count": len(item.sources),
        }

        for key in dimension_keys:
            row[f"score_{key}"] = item.dimension_scores.get(key, 0.0)

        rows.append(row)

    return rows


def _iter_source_lines(sources: Iterable) -> list[str]:
    """
    Genera líneas legibles para fuentes en Markdown.
    """
    lines: list[str] = []
    for source in sources:
        title = source.title or "Fuente sin título"
        url = source.url or "URL no disponible"
        summary = source.summary or ""
        category = source.category or ""

        if category:
            lines.append(f"- **[{category.upper()}] {title}** — {url}")
        else:
            lines.append(f"- **{title}** — {url}")

        if summary:
            lines.append(f"  - {summary}")

    return lines


def build_ranking_markdown_report(ranking_result: RankingResult) -> str:
    """
    Construye un informe Markdown completo del ranking.
    """
    md: List[str] = []

    md.append("# Ranking de Riesgo País")
    md.append("")
    md.append("## Metadatos")
    md.append(f"- Fecha de generación: {ranking_result.metadata.generated_at}")
    md.append(f"- Sector: {ranking_result.metadata.sector}")
    md.append(f"- Tipo de empresa: {ranking_result.metadata.company_type}")
    md.append(f"- Total de países analizados: {ranking_result.metadata.total_countries}")
    md.append(f"- Run ID: {ranking_result.metadata.run_id}")
    md.append("")

    md.append("## Ranking final")
    md.append("")

    for item in ranking_result.ranking:
        md.append(f"### {item.position}. {item.country}")
        md.append(f"- Score total: {item.score_total}")
        md.append("")

        md.append("#### Desglose por dimensiones")
        if item.dimension_scores:
            for key, value in item.dimension_scores.items():
                md.append(f"- {prettify_dimension_name(key)}: {value}")
        else:
            md.append("- Sin dimensiones registradas.")
        md.append("")

        md.append("#### Resumen ejecutivo")
        md.append(item.executive_summary or "Sin resumen disponible.")
        md.append("")

        md.append("#### Fuentes")
        if item.sources:
            md.extend(_iter_source_lines(item.sources))
        else:
            md.append("- No se registraron fuentes.")
        md.append("")

    return "\n".join(md)


def export_ranking_to_json(ranking_result: RankingResult) -> Path:
    """
    Exporta el ranking completo a JSON.
    """
    file_path = build_ranking_filename(
        sector=ranking_result.metadata.sector,
        company_type=ranking_result.metadata.company_type,
        extension="json",
        run_id=ranking_result.metadata.run_id,
    )

    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(ranking_result.model_dump(), f, ensure_ascii=False, indent=4)

    return file_path


def export_ranking_to_csv(ranking_result: RankingResult) -> Path:
    """
    Exporta el ranking a CSV en formato plano.
    """
    file_path = build_ranking_filename(
        sector=ranking_result.metadata.sector,
        company_type=ranking_result.metadata.company_type,
        extension="csv",
        run_id=ranking_result.metadata.run_id,
    )

    rows = flatten_ranking_for_csv(ranking_result)

    if not rows:
        raise ValueError("No hay filas de ranking para exportar a CSV.")

    with open(file_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)

    return file_path


def export_ranking_to_markdown(ranking_result: RankingResult) -> Path:
    """
    Exporta el ranking completo a Markdown.
    """
    file_path = build_ranking_filename(
        sector=ranking_result.metadata.sector,
        company_type=ranking_result.metadata.company_type,
        extension="md",
        run_id=ranking_result.metadata.run_id,
    )

    content = build_ranking_markdown_report(ranking_result)

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)

    return file_path


def export_ranking_to_docx(ranking_result: RankingResult) -> Path:
    """
    Exporta el ranking completo a DOCX.
    """
    file_path = build_ranking_filename(
        sector=ranking_result.metadata.sector,
        company_type=ranking_result.metadata.company_type,
        extension="docx",
        run_id=ranking_result.metadata.run_id,
    )

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    document.add_heading("Ranking de Riesgo País", level=0)
    document.add_paragraph(f"Fecha de generación: {ranking_result.metadata.generated_at}")
    document.add_paragraph(f"Sector: {ranking_result.metadata.sector}")
    document.add_paragraph(f"Tipo de empresa: {ranking_result.metadata.company_type}")
    document.add_paragraph(f"Total de países analizados: {ranking_result.metadata.total_countries}")
    document.add_paragraph(f"Run ID: {ranking_result.metadata.run_id}")

    document.add_heading("Tabla resumen", level=1)
    table = document.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Posición"
    hdr[1].text = "País"
    hdr[2].text = "Score total"

    for item in ranking_result.ranking:
        row = table.add_row().cells
        row[0].text = str(item.position)
        row[1].text = item.country
        row[2].text = str(item.score_total)

    for item in ranking_result.ranking:
        document.add_page_break()
        document.add_heading(f"{item.position}. {item.country}", level=1)
        document.add_paragraph(f"Score total: {item.score_total}")

        document.add_heading("Desglose por dimensiones", level=2)
        if item.dimension_scores:
            for key, value in item.dimension_scores.items():
                document.add_paragraph(f"{prettify_dimension_name(key)}: {value}")
        else:
            document.add_paragraph("Sin dimensiones registradas.")

        document.add_heading("Resumen ejecutivo", level=2)
        document.add_paragraph(item.executive_summary or "Sin resumen disponible.")

        document.add_heading("Fuentes", level=2)
        if item.sources:
            for source in item.sources:
                p = document.add_paragraph(style="List Bullet")
                source_title = source.title or "Fuente sin título"
                source_category = f"[{source.category.upper()}] " if source.category else ""
                p.add_run(f"{source_category}{source_title}")
                if source.url:
                    p.add_run(f" — {source.url}")
                if source.summary:
                    document.add_paragraph(source.summary)
        else:
            document.add_paragraph("No se registraron fuentes.")

    document.save(file_path)
    return file_path


# -------------------------------------------------------------------
# FASE 16 — EXPORTACIÓN EJECUTIVA AVANZADA
# -------------------------------------------------------------------

def _add_portada(
    document: Document,
    titulo: str,
    subtitulo: str,
    sector: str,
    tipo_empresa: str,
    fecha: str,
) -> None:
    """
    Añade una portada profesional al documento DOCX.

    La portada incluye título principal, subtítulo, metadatos
    clave y una línea separadora visual.
    """
    # ---------------------------------------------------------------
    # Título principal — centrado y en negrita
    # ---------------------------------------------------------------
    p_titulo = document.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(20)
    run_titulo.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # ---------------------------------------------------------------
    # Subtítulo — centrado
    # ---------------------------------------------------------------
    p_sub = document.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(subtitulo)
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    document.add_paragraph()

    # ---------------------------------------------------------------
    # Metadatos de portada
    # ---------------------------------------------------------------
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Sector: {sector}  ·  Tipo de empresa: {tipo_empresa}  ·  Fecha: {fecha}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    document.add_paragraph()

    # ---------------------------------------------------------------
    # Línea separadora visual — tabla de 1 fila y 1 celda con borde
    # ---------------------------------------------------------------
    sep = document.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = sep.add_run("─" * 60)
    run_sep.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run_sep.font.size = Pt(9)

    document.add_paragraph()

    # ---------------------------------------------------------------
    # Aviso metodológico
    # ---------------------------------------------------------------
    aviso = document.add_paragraph()
    aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_aviso = aviso.add_run(
        "Informe generado mediante análisis LLM-based con fuentes web en tiempo real. "
        "Los scores son orientativos y no sustituyen una due diligence especializada."
    )
    run_aviso.font.size = Pt(8)
    run_aviso.font.italic = True
    run_aviso.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    document.add_page_break()


def _add_resumen_ejecutivo_consolidado(
    document: Document,
    pais_a: str,
    resultado_a: dict,
    pais_b: str,
    resultado_b: dict,
    comparacion_narrativa: str,
) -> None:
    """
    Añade una sección de resumen ejecutivo consolidado al inicio del informe.

    Esta sección condensa los hallazgos clave de ambos países en una
    página, pensada para lectores que solo leerán el principio del informe.
    """
    document.add_heading("Resumen ejecutivo", level=1)

    # ---------------------------------------------------------------
    # Resumen de cada país en un párrafo corto
    # ---------------------------------------------------------------
    for pais, resultado in [(pais_a, resultado_a), (pais_b, resultado_b)]:
        p = document.add_paragraph()
        run_pais = p.add_run(f"{pais}: ")
        run_pais.bold = True
        run_pais.font.size = Pt(10)
        p.add_run(resultado.get("resumen_ejecutivo", "Sin resumen disponible."))

    document.add_paragraph()

    # ---------------------------------------------------------------
    # Tabla comparativa de scores clave
    # ---------------------------------------------------------------
    document.add_heading("Comparativa de scores clave", level=2)

    scores_a = resultado_a.get("scores", {})
    scores_b = resultado_b.get("scores", {})

    dimensiones_clave = [
        "riesgo_politico",
        "estabilidad_economica",
        "riesgo_regulatorio",
        "riesgo_comercial",
    ]

    tabla = document.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tabla.rows[0].cells
    hdr[0].text = "Dimensión"
    hdr[1].text = pais_a
    hdr[2].text = pais_b

    for dim in dimensiones_clave:
        fila = tabla.add_row().cells
        fila[0].text = prettify_dimension_name(dim)
        fila[1].text = str(scores_a.get(dim, "—"))
        fila[2].text = str(scores_b.get(dim, "—"))

    document.add_paragraph()

    # ---------------------------------------------------------------
    # Conclusión comparativa narrativa
    # ---------------------------------------------------------------
    document.add_heading("Conclusión comparativa", level=2)
    document.add_paragraph(comparacion_narrativa or "Sin conclusión narrativa disponible.")

    document.add_page_break()


def export_to_docx_ejecutivo(
    pais_a: str,
    resultado_a: dict,
    pais_b: str,
    resultado_b: dict,
    comparativa_scores: dict,
    comparacion_narrativa: str,
    sector: str,
    tipo_empresa: str,
    filename_without_ext: str | Path,
) -> Path:
    """
    Exporta la comparación entre dos países a DOCX en versión ejecutiva.

    Fase 16: versión corta pensada para directivos y clientes.
    Incluye portada profesional, resumen ejecutivo consolidado,
    tabla comparativa de scores clave y conclusión narrativa.

    No incluye justificaciones técnicas ni fuentes completas.
    """
    ensure_output_dirs()
    base_path = _normalize_base_path(filename_without_ext)
    output_path = base_path.with_suffix("").with_suffix("") if str(base_path).endswith(".ejecutivo") \
        else Path(str(base_path) + "_ejecutivo.docx")

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ---------------------------------------------------------------
    # Portada profesional
    # ---------------------------------------------------------------
    _add_portada(
        document=document,
        titulo="Informe de Inteligencia Comercial",
        subtitulo=f"Análisis comparativo: {pais_a} vs {pais_b}",
        sector=sector,
        tipo_empresa=tipo_empresa,
        fecha=datetime.now().strftime("%d/%m/%Y"),
    )

    # ---------------------------------------------------------------
    # Resumen ejecutivo consolidado con tabla comparativa
    # ---------------------------------------------------------------
    _add_resumen_ejecutivo_consolidado(
        document=document,
        pais_a=pais_a,
        resultado_a=resultado_a,
        pais_b=pais_b,
        resultado_b=resultado_b,
        comparacion_narrativa=comparacion_narrativa,
    )

    # ---------------------------------------------------------------
    # Alertas y oportunidades por país
    # ---------------------------------------------------------------
    document.add_heading("Alertas y oportunidades", level=1)

    for pais, resultado in [(pais_a, resultado_a), (pais_b, resultado_b)]:
        document.add_heading(pais, level=2)

        alertas = resultado.get("alertas", [])
        document.add_heading("Alertas", level=3)
        if alertas:
            for alerta in alertas:
                p = document.add_paragraph(style="List Bullet")
                p.add_run(alerta)
        else:
            document.add_paragraph("Sin alertas registradas.")

        oportunidades = resultado.get("oportunidades", [])
        document.add_heading("Oportunidades", level=3)
        if oportunidades:
            for oportunidad in oportunidades:
                p = document.add_paragraph(style="List Bullet")
                p.add_run(oportunidad)
        else:
            document.add_paragraph("Sin oportunidades registradas.")

    document.save(output_path)
    return output_path


def export_to_docx_tecnico(
    pais_a: str,
    resultado_a: dict,
    pais_b: str,
    resultado_b: dict,
    comparativa_scores: dict,
    comparacion_narrativa: str,
    sector: str,
    tipo_empresa: str,
    filename_without_ext: str | Path,
) -> Path:
    """
    Exporta la comparación entre dos países a DOCX en versión técnica ampliada.

    Fase 16: versión completa pensada para analistas y equipos técnicos.
    Incluye portada, resumen ejecutivo, scores completos con justificaciones
    por dimensión, bandas de confianza y fuentes trazables completas.
    """
    ensure_output_dirs()
    base_path = _normalize_base_path(filename_without_ext)
    output_path = Path(str(base_path) + "_tecnico.docx")

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ---------------------------------------------------------------
    # Portada profesional
    # ---------------------------------------------------------------
    _add_portada(
        document=document,
        titulo="Informe Técnico de Inteligencia Comercial",
        subtitulo=f"Análisis comparativo ampliado: {pais_a} vs {pais_b}",
        sector=sector,
        tipo_empresa=tipo_empresa,
        fecha=datetime.now().strftime("%d/%m/%Y"),
    )

    # ---------------------------------------------------------------
    # Resumen ejecutivo consolidado
    # ---------------------------------------------------------------
    _add_resumen_ejecutivo_consolidado(
        document=document,
        pais_a=pais_a,
        resultado_a=resultado_a,
        pais_b=pais_b,
        resultado_b=resultado_b,
        comparacion_narrativa=comparacion_narrativa,
    )

    # ---------------------------------------------------------------
    # Análisis técnico completo por país
    # ---------------------------------------------------------------
    document.add_heading("Análisis técnico por país", level=1)

    for pais, resultado in [(pais_a, resultado_a), (pais_b, resultado_b)]:
        document.add_page_break()
        document.add_heading(pais, level=2)

        # Scores completos con bandas de confianza
        document.add_heading("Scores por dimensión", level=3)
        scores = resultado.get("scores", {})
        bandas = resultado.get("bandas_confianza", {})

        if scores:
            tabla_scores = document.add_table(rows=1, cols=3)
            tabla_scores.style = "Table Grid"
            hdr = tabla_scores.rows[0].cells
            hdr[0].text = "Dimensión"
            hdr[1].text = "Score (1–10)"
            hdr[2].text = "Confianza"

            for dim, valor in scores.items():
                fila = tabla_scores.add_row().cells
                fila[0].text = prettify_dimension_name(dim)
                fila[1].text = str(valor)
                fila[2].text = bandas.get(dim, "—")
        else:
            document.add_paragraph("Sin scores registrados.")

        document.add_paragraph()

        # Justificaciones técnicas por dimensión
        document.add_heading("Justificación de scores", level=3)
        justificaciones = resultado.get("justificacion_scores", {})
        if justificaciones:
            for dimension, motivos in justificaciones.items():
                p_dim = document.add_paragraph()
                run_dim = p_dim.add_run(prettify_dimension_name(dimension))
                run_dim.bold = True
                run_dim.font.size = Pt(10)
                for motivo in motivos:
                    p = document.add_paragraph(style="List Bullet")
                    p.add_run(motivo)
        else:
            document.add_paragraph("Sin justificaciones registradas.")

        # Alertas y oportunidades
        document.add_heading("Alertas", level=3)
        alertas = resultado.get("alertas", [])
        if alertas:
            for alerta in alertas:
                p = document.add_paragraph(style="List Bullet")
                p.add_run(alerta)
        else:
            document.add_paragraph("Sin alertas registradas.")

        document.add_heading("Oportunidades", level=3)
        oportunidades = resultado.get("oportunidades", [])
        if oportunidades:
            for oportunidad in oportunidades:
                p = document.add_paragraph(style="List Bullet")
                p.add_run(oportunidad)
        else:
            document.add_paragraph("Sin oportunidades registradas.")

        # Fuentes completas con trazabilidad
        document.add_heading("Fuentes utilizadas", level=3)
        fuentes = resultado.get("fuentes", [])
        if fuentes:
            for fuente in fuentes:
                categoria = fuente.get("categoria", "sin_categoria")
                titulo = fuente.get("titulo", "Sin título")
                url = fuente.get("url", "")
                resumen_fuente = fuente.get("resumen", "")

                p = document.add_paragraph(style="List Bullet")
                run_cat = p.add_run(f"[{categoria.upper()}] ")
                run_cat.bold = True
                p.add_run(titulo)
                if url:
                    p.add_run(f" — {url}")
                if resumen_fuente:
                    p_res = document.add_paragraph(resumen_fuente)
                    p_res.runs[0].font.size = Pt(9)
                    p_res.runs[0].font.italic = True
        else:
            document.add_paragraph("No se registraron fuentes.")

    # ---------------------------------------------------------------
    # Sección de scoring agregado comparativo
    # ---------------------------------------------------------------
    document.add_page_break()
    document.add_heading("Scoring agregado comparativo", level=1)

    scores_a = resultado_a.get("scores", {})
    scores_b = resultado_b.get("scores", {})
    todas_dimensiones = sorted(set(list(scores_a.keys()) + list(scores_b.keys())))

    if todas_dimensiones:
        tabla_comp = document.add_table(rows=1, cols=3)
        tabla_comp.style = "Table Grid"
        tabla_comp.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = tabla_comp.rows[0].cells
        hdr[0].text = "Dimensión"
        hdr[1].text = pais_a
        hdr[2].text = pais_b

        for dim in todas_dimensiones:
            fila = tabla_comp.add_row().cells
            fila[0].text = prettify_dimension_name(dim)
            fila[1].text = str(scores_a.get(dim, "—"))
            fila[2].text = str(scores_b.get(dim, "—"))

    document.add_paragraph()
    document.add_heading("Metodología", level=2)
    document.add_paragraph(
        "El scoring se calcula mediante análisis semántico LLM-based sobre contexto web "
        "recuperado en tiempo real por Tavily. Cada dimensión recibe un score entre 1 (muy favorable) "
        "y 10 (muy desfavorable), con justificación explícita y banda de confianza según la "
        "cobertura de evidencia detectada (alta ≥3 señales · media 1-2 señales · baja sin señales)."
    )

    document.save(output_path)
    return output_path


def export_ranking_to_docx_ejecutivo(ranking_result: RankingResult) -> Path:
    """
    Exporta el ranking a DOCX en versión ejecutiva con portada.

    Fase 16: incluye portada, tabla resumen y resúmenes ejecutivos
    del top 3. Pensado para presentaciones a dirección.
    """
    file_path = build_ranking_filename(
        sector=ranking_result.metadata.sector,
        company_type=ranking_result.metadata.company_type,
        extension="docx",
        prefix="ranking_ejecutivo",
        run_id=ranking_result.metadata.run_id,
    )

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ---------------------------------------------------------------
    # Portada
    # ---------------------------------------------------------------
    _add_portada(
        document=document,
        titulo="Ranking de Inteligencia Comercial",
        subtitulo=f"{ranking_result.metadata.total_countries} mercados analizados",
        sector=ranking_result.metadata.sector,
        tipo_empresa=ranking_result.metadata.company_type,
        fecha=ranking_result.metadata.generated_at[:10],
    )

    # ---------------------------------------------------------------
    # Tabla resumen de ranking
    # ---------------------------------------------------------------
    document.add_heading("Ranking de mercados", level=1)

    tabla = document.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tabla.rows[0].cells
    hdr[0].text = "Posición"
    hdr[1].text = "País"
    hdr[2].text = "Score total"

    for item in ranking_result.ranking:
        fila = tabla.add_row().cells
        fila[0].text = str(item.position)
        fila[1].text = item.country
        fila[2].text = str(item.score_total)

    document.add_paragraph()

    # ---------------------------------------------------------------
    # Resúmenes ejecutivos del top 3
    # ---------------------------------------------------------------
    document.add_heading("Análisis del Top 3", level=1)

    for item in ranking_result.ranking[:3]:
        document.add_heading(f"{item.position}. {item.country}", level=2)
        document.add_paragraph(f"Score total: {item.score_total}")
        document.add_paragraph(item.executive_summary or "Sin resumen disponible.")
        document.add_paragraph()

    document.save(file_path)
    return file_path 