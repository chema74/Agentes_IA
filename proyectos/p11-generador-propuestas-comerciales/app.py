"""
P23 - Base publica del agente de ventas y estrategia de negocio
================================================================
Autor : Jose Maria
Stack : Groq - python-docx - Streamlit
Coste : GRATUITO
"""

import io
import json
import os
from datetime import datetime

import streamlit as st
from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.shared import Inches, Pt, RGBColor  # type: ignore
from dotenv import load_dotenv
from groq import Groq

load_dotenv()

CLAVES_PROPUESTA = {
    "resumen_ejecutivo",
    "problema",
    "solucion",
    "beneficios",
    "metodologia",
    "por_que_nosotros",
    "inversion",
    "siguiente_paso",
}

st.set_page_config(
    page_title="Agente de ventas y estrategia de negocio",
    page_icon=":page_facing_up:",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,300;0,9..144,700;0,9..144,900;1,9..144,300&family=DM+Mono:wght@400;500&family=DM+Sans:wght@300;400;500&display=swap');
html, body, [class*="css"] { font-family:'DM Sans',sans-serif; background:#0c0c10; color:#e4e2dc; }
.stApp { background:#0c0c10; }
#MainMenu, footer, header { visibility:hidden; }
.block-container { padding-top:2rem; padding-bottom:2rem; max-width:1100px; }
.app-tag { font-family:'DM Mono',monospace; font-size:.65rem; letter-spacing:.2em; text-transform:uppercase; color:#d4a84b; margin-bottom:.75rem; }
.app-title { font-family:'Fraunces',serif; font-size:2.2rem; font-weight:900; line-height:1.1; margin:0; }
.app-title em { font-style:italic; font-weight:300; color:#d4a84b; }
.app-subtitle { color:#8c8a84; font-size:.9rem; margin-top:.5rem; }
.app-header { border-bottom:1px solid rgba(212,168,75,.2); padding-bottom:1.5rem; margin-bottom:2rem; }
.groq-badge { display:inline-flex; align-items:center; gap:.4rem; font-family:'DM Mono',monospace; font-size:.6rem; color:#4dd488; border:1px solid rgba(77,212,136,.25); padding:.2rem .6rem; margin-left:.75rem; }
.stTextInput > label, .stTextArea > label, .stSelectbox > label { font-family:'DM Mono',monospace !important; font-size:.7rem !important; letter-spacing:.12em !important; text-transform:uppercase !important; color:#d4a84b !important; }
.stTextInput input, .stTextArea textarea { background:#14141c !important; border:1px solid rgba(212,168,75,.25) !important; border-radius:3px !important; color:#e4e2dc !important; }
[data-baseweb="select"] > div { background:#14141c !important; border:1px solid rgba(212,168,75,.25) !important; border-radius:3px !important; color:#e4e2dc !important; }
.stButton > button { background:#d4a84b !important; color:#0c0c10 !important; border:none !important; border-radius:3px !important; font-family:'DM Mono',monospace !important; font-size:.75rem !important; font-weight:700 !important; letter-spacing:.1em !important; text-transform:uppercase !important; padding:.65rem 2rem !important; }
.stButton > button:hover { background:#e8c97a !important; transform:translateY(-1px); }
.preview-block { background:#14141c; border:1px solid rgba(212,168,75,.15); padding:1.5rem 2rem; margin-bottom:1rem; position:relative; }
.preview-block::before { content:''; position:absolute; top:0; left:0; width:3px; height:100%; background:linear-gradient(180deg,#d4a84b,transparent); }
.preview-label { font-family:'DM Mono',monospace; font-size:.62rem; letter-spacing:.12em; text-transform:uppercase; color:#7a5e28; margin-bottom:.6rem; }
.preview-text { font-size:.9rem; color:#e4e2dc; line-height:1.85; }
.custom-divider { height:1px; background:linear-gradient(90deg,transparent,rgba(212,168,75,.3),transparent); margin:1.5rem 0; }
.app-footer { font-family:'DM Mono',monospace; font-size:.62rem; color:#44433f; text-align:center; padding-top:2rem; }
[data-testid="stSidebar"] { background:#10101a !important; border-right:1px solid rgba(212,168,75,.12) !important; }
</style>
""",
    unsafe_allow_html=True,
)


@st.cache_resource
def get_groq():
    api_key = os.getenv("GROQ_API_KEY", "").strip()
    if not api_key:
        raise ValueError(
            "Falta GROQ_API_KEY. Configura tu clave en el archivo .env para usar el generador."
        )
    return Groq(api_key=api_key)


def extraer_objetos_json(texto: str):
    """Extrae objetos JSON balanceados sin depender de find/rfind."""
    objetos = []
    inicio = None
    profundidad = 0
    en_cadena = False
    escape = False

    for i, ch in enumerate(texto):
        if en_cadena:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                en_cadena = False
            continue

        if ch == '"':
            en_cadena = True
        elif ch == "{":
            if profundidad == 0:
                inicio = i
            profundidad += 1
        elif ch == "}" and profundidad > 0:
            profundidad -= 1
            if profundidad == 0 and inicio is not None:
                objetos.append(texto[inicio : i + 1])
                inicio = None

    return objetos


def validar_propuesta(data: dict) -> dict:
    if not isinstance(data, dict):
        raise ValueError("La respuesta del modelo no tiene formato JSON de objeto.")

    faltantes = sorted(CLAVES_PROPUESTA - set(data.keys()))
    if faltantes:
        raise ValueError(
            "La propuesta generada esta incompleta. Faltan secciones: "
            + ", ".join(faltantes)
        )

    for campo in ["beneficios", "metodologia"]:
        if not isinstance(data.get(campo), list):
            raise ValueError(f"El campo {campo} debe ser una lista.")

    return data


def parsear_respuesta_json(raw: str) -> dict:
    if not raw or not raw.strip():
        raise ValueError("Groq devolvio una respuesta vacia.")

    bloques = [raw.strip()]
    if "```" in raw:
        for parte in raw.split("```"):
            parte = parte.strip()
            if parte.lower().startswith("json"):
                parte = parte[4:].strip()
            if parte:
                bloques.append(parte)

    errores = []
    vistos = set()
    for bloque in bloques:
        objetos = sorted(extraer_objetos_json(bloque), key=len, reverse=True)
        for candidato in objetos:
            if candidato in vistos:
                continue
            vistos.add(candidato)
            try:
                data = json.loads(candidato)
                return validar_propuesta(data)
            except (json.JSONDecodeError, ValueError) as ex:
                errores.append(str(ex))

    if errores:
        raise ValueError("No se pudo validar la propuesta generada. " + errores[0])
    raise ValueError("No se encontro un JSON valido en la respuesta del modelo.")


def generar_propuesta(groq_client, datos: dict) -> dict:
    prompt = f"""Genera una propuesta comercial profesional en espanol con los siguientes datos:

EMPRESA EMISORA: {datos['empresa_emisora']}
CLIENTE: {datos['nombre_cliente']} ({datos['empresa_cliente']})
SECTOR DEL CLIENTE: {datos['sector_cliente']}
NECESIDAD DETECTADA: {datos['necesidad']}
SOLUCION A OFRECER: {datos['solucion']}
PRECIO / INVERSION: {datos['precio']}
PLAZO DE ENTREGA: {datos['plazo']}
CONTACTO: {datos['contacto_nombre']} - {datos['contacto_email']}
TONO DESEADO: {datos['tono']}

Genera un JSON con exactamente estas secciones:
resumen_ejecutivo: 3-4 frases que resuman el valor de la propuesta
problema: 2-3 frases describiendo el problema o necesidad del cliente
solucion: descripcion de la solucion propuesta en 3-4 frases, orientada a beneficios
beneficios: lista de 4-5 beneficios concretos para el cliente
metodologia: lista de 4-5 pasos del proceso de trabajo o implementacion
por_que_nosotros: 3-4 frases sobre por que el cliente deberia elegir esta empresa
inversion: descripcion del precio y condiciones de pago
siguiente_paso: 1-2 frases con la llamada a la accion clara y concreta

REGLAS:
- Usa el tono indicado en TONO DESEADO.
- Lenguaje claro, profesional y orientado a negocio.
- No inventes capacidades no descritas.
- Responde solo con JSON valido. Sin markdown."""

    response = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=1500,
    )
    raw = (response.choices[0].message.content or "").strip()
    return parsear_respuesta_json(raw)


def crear_word(propuesta: dict, datos: dict) -> bytes:
    """Genera el documento Word de la propuesta."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    amber = RGBColor(0xC9, 0x95, 0x30)
    dark = RGBColor(0x1A, 0x1A, 0x26)
    gray = RGBColor(0x55, 0x55, 0x55)

    def add_heading_custom(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(22)
            run.font.color.rgb = dark
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = amber
            run.font.name = "Courier New"
        return p

    def add_body(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = gray
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = gray
        return p

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("-" * 70)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0xCC, 0xAA, 0x44)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(datos["empresa_emisora"].upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = amber
    run.font.name = "Courier New"

    doc.add_paragraph()
    add_heading_custom(f"Propuesta para\n{datos['empresa_cliente']}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(datos["solucion"][:60] + "...")
    run.font.size = Pt(12)
    run.font.color.rgb = gray
    run.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        f"Preparada para: {datos['nombre_cliente']} - {datetime.now().strftime('%d/%m/%Y')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = "Courier New"

    add_divider()
    doc.add_page_break()

    secciones = [
        ("// Resumen ejecutivo", propuesta.get("resumen_ejecutivo", ""), "text"),
        ("// El desafio", propuesta.get("problema", ""), "text"),
        ("// Nuestra solucion", propuesta.get("solucion", ""), "text"),
        ("// Beneficios para el cliente", propuesta.get("beneficios", []), "list"),
        ("// Metodologia de trabajo", propuesta.get("metodologia", []), "list"),
        ("// Por que elegirnos", propuesta.get("por_que_nosotros", ""), "text"),
        ("// Inversion", propuesta.get("inversion", ""), "text"),
        ("// Siguiente paso", propuesta.get("siguiente_paso", ""), "text"),
    ]

    for titulo, contenido, tipo in secciones:
        add_heading_custom(titulo, level=2)
        doc.add_paragraph()
        if tipo == "text":
            add_body(str(contenido))
        else:
            for item in (contenido if isinstance(contenido, list) else [contenido]):
                add_bullet(str(item))
        doc.add_paragraph()

    add_divider()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{datos['contacto_nombre']} - {datos['contacto_email']} - {datos['empresa_emisora']}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = gray
    run.font.name = "Courier New"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


if "propuesta_generada" not in st.session_state:
    st.session_state.propuesta_generada = None
if "datos_propuesta" not in st.session_state:
    st.session_state.datos_propuesta = None


with st.sidebar:
    st.markdown(
        '<div style="font-family:\'DM Mono\',monospace;font-size:.65rem;letter-spacing:.15em;text-transform:uppercase;color:#d4a84b;margin-bottom:1.25rem">// Tu empresa</div>',
        unsafe_allow_html=True,
    )
    empresa_emisora = st.text_input(
        "Nombre de tu empresa", placeholder="Ej: Consulting IA S.L."
    )
    contacto_nombre = st.text_input(
        "Tu nombre", placeholder="Ej: Jose Maria Garcia"
    )
    contacto_email = st.text_input("Tu email", placeholder="tu@empresa.com")

    st.markdown(
        '<div style="font-family:\'DM Mono\',monospace;font-size:.65rem;letter-spacing:.15em;text-transform:uppercase;color:#d4a84b;margin-bottom:.75rem;margin-top:1rem">// Cliente</div>',
        unsafe_allow_html=True,
    )
    nombre_cliente = st.text_input(
        "Nombre del contacto", placeholder="Ej: Ana Martinez"
    )
    empresa_cliente = st.text_input(
        "Empresa del cliente", placeholder="Ej: Distribuciones Garcia S.A."
    )
    sector_cliente = st.text_input(
        "Sector", placeholder="Ej: Distribucion alimentaria"
    )

    st.markdown(
        """
    <div style="font-family:'DM Mono',monospace;font-size:.6rem;color:#44433f;line-height:1.9;border-top:1px solid rgba(212,168,75,.1);padding-top:1rem;margin-top:1rem">
        - Genera propuestas comerciales con Groq<br>
        - Exporta a Word (.docx)<br>
        - Base publica actual del agente de ventas y estrategia
    </div>""",
        unsafe_allow_html=True,
    )


st.markdown(
    """
<div class="app-header">
  <div class="app-tag">Agente de ventas y estrategia de negocio
    <span class="groq-badge">Base publica actual - Groq - Word</span>
  </div>
  <div class="app-title">Propuestas <em>listas para enviar</em></div>
  <div class="app-subtitle">Version publica actual: genera propuestas comerciales entregables a partir de informacion comercial estructurada</div>
</div>""",
    unsafe_allow_html=True,
)

col1, col2 = st.columns(2)
with col1:
    necesidad = st.text_area(
        "Necesidad o problema detectado en el cliente",
        placeholder="Ej: La empresa no tiene sistematizado el seguimiento de sus leads y pierde oportunidades por falta de seguimiento.",
        height=100,
    )
    solucion = st.text_area(
        "Solucion que propones",
        placeholder="Ej: Implementacion de un CRM con IA para automatizar seguimiento, priorizar oportunidades y mejorar conversion.",
        height=100,
    )
with col2:
    precio = st.text_input(
        "Precio / inversion", placeholder="Ej: 2.400 EUR + IVA en 2 plazos"
    )
    plazo = st.text_input(
        "Plazo de entrega o implementacion",
        placeholder="Ej: 4 semanas desde la firma",
    )
    tono = st.selectbox(
        "Tono de la propuesta",
        ["Profesional y cercano", "Formal y tecnico", "Directo y ejecutivo"],
    )

st.markdown("<div style='height:.5rem'></div>", unsafe_allow_html=True)

col_btn, col_info = st.columns([1, 3])
with col_btn:
    generar_btn = st.button("Generar propuesta", use_container_width=True)
with col_info:
    st.markdown(
        '<div style="font-family:\'DM Mono\',monospace;font-size:.62rem;color:#44433f;padding:.75rem 0">Genera un borrador comercial entregable y descargalo en Word para revision final.</div>',
        unsafe_allow_html=True,
    )

if generar_btn:
    campos_requeridos = [
        empresa_emisora,
        contacto_nombre,
        contacto_email,
        nombre_cliente,
        empresa_cliente,
        necesidad,
        solucion,
        precio,
    ]
    if not all(campos_requeridos):
        st.warning(
            "Rellena al menos: tu empresa, tu nombre, email, nombre del cliente, empresa del cliente, necesidad, solucion y precio."
        )
    else:
        datos = {
            "empresa_emisora": empresa_emisora,
            "contacto_nombre": contacto_nombre,
            "contacto_email": contacto_email,
            "nombre_cliente": nombre_cliente,
            "empresa_cliente": empresa_cliente,
            "sector_cliente": sector_cliente or "No especificado",
            "necesidad": necesidad,
            "solucion": solucion,
            "precio": precio,
            "plazo": plazo or "A convenir",
            "tono": tono,
        }
        with st.spinner("Generando propuesta con Groq..."):
            try:
                propuesta = generar_propuesta(get_groq(), datos)
                st.session_state.propuesta_generada = propuesta
                st.session_state.datos_propuesta = datos
                st.rerun()
            except Exception as ex:
                st.error(f"No se pudo generar la propuesta: {ex}")


if st.session_state.propuesta_generada and st.session_state.datos_propuesta:
    p = st.session_state.propuesta_generada
    d = st.session_state.datos_propuesta

    st.markdown(
        f"""
    <div style="display:flex;justify-content:space-between;align-items:baseline;
        border-bottom:1px solid rgba(212,168,75,.2);padding-bottom:1rem;margin-bottom:1.5rem">
      <div>
        <span style="font-family:'Fraunces',serif;font-size:1.4rem;font-weight:700">
          Propuesta para {d['empresa_cliente']}
        </span>
      </div>
      <div style="font-family:'DM Mono',monospace;font-size:.6rem;color:#44433f">
        {datetime.now().strftime('%d/%m/%Y')} - {d['empresa_emisora']}
      </div>
    </div>""",
        unsafe_allow_html=True,
    )

    secciones_preview = [
        ("Resumen ejecutivo", p.get("resumen_ejecutivo", "")),
        ("El desafio", p.get("problema", "")),
        ("Nuestra solucion", p.get("solucion", "")),
        ("Por que elegirnos", p.get("por_que_nosotros", "")),
        ("Inversion", p.get("inversion", "")),
        ("Siguiente paso", p.get("siguiente_paso", "")),
    ]

    col_a, col_b = st.columns(2)
    for i, (titulo, contenido) in enumerate(secciones_preview):
        with (col_a if i % 2 == 0 else col_b):
            contenido_str = (
                "<br>".join([f"- {item}" for item in contenido])
                if isinstance(contenido, list)
                else str(contenido)
            )
            st.markdown(
                f"""
            <div class="preview-block">
              <div class="preview-label">{titulo}</div>
              <div class="preview-text">{contenido_str}</div>
            </div>""",
                unsafe_allow_html=True,
            )

    st.markdown("<div class='custom-divider'></div>", unsafe_allow_html=True)
    col_c, col_d = st.columns(2)
    with col_c:
        st.markdown(
            '<div style="font-family:\'DM Mono\',monospace;font-size:.62rem;color:#7a5e28;text-transform:uppercase;letter-spacing:.1em;margin-bottom:.6rem">Beneficios para el cliente</div>',
            unsafe_allow_html=True,
        )
        for beneficio in p.get("beneficios", []):
            st.markdown(
                f'<div style="padding:.4rem 0;border-bottom:1px solid rgba(212,168,75,.07);font-size:.875rem">- {beneficio}</div>',
                unsafe_allow_html=True,
            )
    with col_d:
        st.markdown(
            '<div style="font-family:\'DM Mono\',monospace;font-size:.62rem;color:#7a5e28;text-transform:uppercase;letter-spacing:.1em;margin-bottom:.6rem">Metodologia de trabajo</div>',
            unsafe_allow_html=True,
        )
        for i, paso in enumerate(p.get("metodologia", []), 1):
            st.markdown(
                f'<div style="padding:.4rem 0;border-bottom:1px solid rgba(212,168,75,.07);font-size:.875rem">{i}. {paso}</div>',
                unsafe_allow_html=True,
            )

    st.markdown("<div class='custom-divider'></div>", unsafe_allow_html=True)

    col_dl, col_new = st.columns(2)
    with col_dl:
        try:
            word_bytes = crear_word(p, d)
            filename = (
                f"propuesta_{d['empresa_cliente'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
            )
            st.download_button(
                "Descargar propuesta en Word",
                data=word_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as ex:
            st.error(f"No se pudo generar el archivo Word: {ex}")
    with col_new:
        if st.button("Nueva propuesta", use_container_width=True):
            st.session_state.propuesta_generada = None
            st.session_state.datos_propuesta = None
            st.rerun()
else:
    st.markdown(
        """
    <div style="border:1px dashed rgba(212,168,75,.2);padding:3rem 2rem;text-align:center;margin-top:1rem">
      <div style="font-size:2.5rem;margin-bottom:1rem">WORD</div>
      <div style="font-family:'Fraunces',serif;font-size:1.2rem;color:#8c8a84;margin-bottom:.75rem">
        Genera tu propuesta comercial y dejala lista para revision final
      </div>
      <div style="font-family:'DM Mono',monospace;font-size:.63rem;color:#44433f;letter-spacing:.06em;line-height:2">
        Esta version publica cubre la generacion de propuestas comerciales entregables<br>
        Actua como base actual del agente de ventas y estrategia de negocio<br>
        <span style="color:#4dd488">Borrador profesional en segundos para personalizar y enviar</span>
      </div>
    </div>""",
        unsafe_allow_html=True,
    )

st.markdown("<div style='height:2rem'></div>", unsafe_allow_html=True)
st.markdown(
    '<div class="app-footer">Agente de ventas y estrategia de negocio - Base publica actual - Groq + python-docx - Portfolio IA Aplicada</div>',
    unsafe_allow_html=True,
)
